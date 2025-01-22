from PIL import Image
from docx import Document
import os
import subprocess 
import argparse

def extract_text_and_media_from_docx(input_file, media_folder):
    """
    Extracts text and images from a DOCX file.

    Args:
        input_file (str): Path to the DOCX file.
        media_folder (str): Path to save extracted images.

    Returns:
        list: A list of AsciiDoc-formatted strings for text and media.
    """

    if not os.path.exists(media_folder): 
        os.makedirs(media_folder) 
    doc = Document(input_file)
    adoc_content = []
    print(doc.paragraphs)

    # read content of docx file and append content to adoc
    for paragraph in doc.paragraphs:
        print(f"Paragraph: {paragraph.text}")
        if paragraph.style.name.startswith("Heading"): #matching on headings"
            level = int(paragraph.style.name.split()[-1])
            adoc_content.append(f"{'=' * level} {paragraph.text}\n")
        else:
            if paragraph.text.strip():
                adoc_content.apppend(prargraph.text + "\n")

    print("Processing tables...")
    for table in doc.tables:
        adoc_content.append("[cols=\"auto\", options=\"header\"]\n|===\n")
        for row in table.rows:
            row_data = "| " + " | ".join(cell.text.strip() for cell in row.cells)
            adoc_content.append(row_data + "\n")
        adoc_content.append("|===\n")
        
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            img_blob = rel.target_part.blob
            img_ext = os.path.splittext(rel.target_ref)[-1]
            img_filename = f"image_{i + 1}{img_ext}"
            img_path = os.path.join(media_folder, img_filename)

            with open(img_path, "wb") as img_file:
                img_file.write(img_blob)
            adoc_content.append(f"image::{img_path}[Image {i + 1}]\n")
    print(adoc_content)
    return adoc_content
        

def convert_images_to_png(media_folder):
    """
    Converts all emf and wmf images in our media folder to png images
    
    Args:
        media_folder (str): Path to the folder
    """
    
    for filename in os.listdir(media_folder):
        if filename.lower().endswith(".emf", ".wmf"):
            file_path = os.path.join(media_folder, filename)
            png_path = os.path.splitext(file_path)[0] + ".png"

            try:
                with Image.open(file_path) as img:
                    img.save(png_path, "PNG")
                    print(f"Converted: {filename} -> {os.path.basename(png_path)}")
                os.remove(file_path)  # Remove the original EMF/WMF file
            except Exception as e:
                print(f"Failed to convert {filename}: {e}")

def write_to_adoc_file(adoc_content, output_file):
    """
    Writes AsciiDoc-formatted content to a file.

    Args:
        adoc_content (list): List of AsciiDoc-formatted strings.
        output_file (str): Path to save the output AsciiDoc file.
    """

    try:
        with open(output_file, "w", encoding="utf-8") as f:
            f.writelines(adoc_content)
        print(f"AsciiDoc file saved to {output_file}")
    except Exception as e:
        print(f"Failed to wrtie AsciiDoc file: {e}")


def main():
    '''
    Main fuction will take two arguments (input_file and output_file) and run a docx to asciidoc conversion of the file, put all media in to its own folder "/media/media", edit the asciidoc file to change all .emf strings to .png strings and convert all emf images to png images"
    
    '''
    parser = argparse.ArgumentParser(description="convert docx to adoc, including image support")
    parser.add_argument("-i", "--input", required=True, help="Docx to convert")
    parser.add_argument("-o", "--output", required=True, help="name of file to convert to")

    args = parser.parse_args()
    media_folder = f"{os.path.splitext(args.output)[0]}_media"

    # Extract text and media
    adoc_content = extract_text_and_media_from_docx(args.input, media_folder)

    # Convert EMF and WMF images to PNG
    convert_images_to_png(media_folder)

    # Write to AsciiDoc file
    write_to_adoc_file(adoc_content, args.output)

if __name__ == "__main__":
    main()